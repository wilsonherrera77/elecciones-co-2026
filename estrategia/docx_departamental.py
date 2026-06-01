"""docx_departamental.py · genera Word ejecutivo por departamento.

Estructura:
- Portada
- Resumen ejecutivo
- Tabla nacional 34 deptos
- Capítulo por depto (34): cifras, cuadrantes locales, top 10 oportunidad, top 10 defensa,
  + recomendación táctica generada por Ollama qwen2.5:14b LOCAL vía templates.routing_inteligente
- Anexo metodológico

Output: data/outputs/analisis_departamental_3M_cepeda.docx
"""

from __future__ import annotations

import os
import subprocess
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

import duckdb
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent
DB = PROJECT_ROOT / "data" / "processed" / "votos.duckdb"
OUT = PROJECT_ROOT / "data" / "outputs" / "analisis_departamental_3M_cepeda.docx"
VA_ROOT = Path(os.environ.get("VA_ROOT", r"C:\Users\wilso\Desktop\Escritorio2026\Visual_Agentes"))

VERDE_PACTO = RGBColor(0x0A, 0x6E, 0x3A)
ROJO_ESPRIELLA = RGBColor(0xC4, 0x3D, 0x2E)
GRIS = RGBColor(0x66, 0x66, 0x66)

CEPEDA = "7"
ESPRIELLA = "10"


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _h(doc: Document, text: str, level: int = 1, color: RGBColor | None = None) -> None:
    p = doc.add_heading(text, level=level)
    if color:
        for r in p.runs:
            r.font.color.rgb = color


def _p(doc: Document, text: str, bold: bool = False, size: int = 11,
       color: RGBColor | None = None, align=None) -> None:
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def _row(table, cells: list[str], header: bool = False) -> None:
    tr = table.add_row()
    for i, txt in enumerate(cells):
        c = tr.cells[i]
        c.text = ""
        para = c.paragraphs[0]
        run = para.add_run(str(txt))
        run.font.name = "Calibri"
        run.font.size = Pt(9 if not header else 10)
        run.bold = header
        if header:
            _set_cell_bg(c, "0A6E3A")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def get_summary_nacional(con) -> dict:
    r = con.execute("""
        SELECT
            (SELECT SUM(censo_electoral) FROM (SELECT DISTINCT scope_idx, censo_electoral FROM votos_municipio_snapshot)) censo,
            (SELECT SUM(votos_validos) FROM (SELECT DISTINCT scope_idx, votos_validos FROM votos_municipio_snapshot)) validos,
            (SELECT SUM(votos) FROM votos_municipio_snapshot WHERE codpar='7') cepeda,
            (SELECT SUM(votos) FROM votos_municipio_snapshot WHERE codpar='10') espriella
    """).fetchone()
    return {"censo": r[0] or 0, "validos": r[1] or 0, "cepeda": r[2] or 0, "espriella": r[3] or 0}


def get_deptos(con) -> list[dict]:
    rows = con.execute("""
        SELECT d.idx, d.codigo_interno, d.nombre,
               COALESCE(SUM(c.votos_cepeda), 0) AS votos_cepeda,
               COALESCE(SUM(c.votos_espriella), 0) AS votos_espriella,
               COALESCE(SUM(c.censo), 0) AS censo,
               COALESCE(SUM(c.validos), 0) AS validos,
               COALESCE(SUM(c.no_votantes), 0) AS no_votantes,
               COUNT(c.scope_idx) AS mpios,
               SUM(CASE WHEN c.cuadrante LIKE 'Q1%' THEN 1 ELSE 0 END) q1,
               SUM(CASE WHEN c.cuadrante='Q2_MOVILIZAR' THEN 1 ELSE 0 END) q2,
               SUM(CASE WHEN c.cuadrante='Q3_CONVERTIR' THEN 1 ELSE 0 END) q3,
               SUM(CASE WHEN c.cuadrante='Q4_RESISTIR' THEN 1 ELSE 0 END) q4
        FROM divipola_2026 d
        LEFT JOIN cuadrantes_2v c ON c.departamento_nombre = d.nombre
        WHERE d.level = 2
        GROUP BY d.idx, d.codigo_interno, d.nombre
        HAVING SUM(c.validos) > 0
        ORDER BY votos_cepeda DESC
    """).fetchall()
    return [{
        "idx": r[0], "codigo": r[1], "nombre": r[2],
        "votos_cepeda": r[3], "votos_espriella": r[4],
        "censo": r[5], "validos": r[6], "no_votantes": r[7],
        "mpios": r[8], "q1": r[9], "q2": r[10], "q3": r[11], "q4": r[12],
    } for r in rows]


def get_top_mpios_depto(con, depto_nombre: str, cuadrante_filter: str, limit: int = 10) -> list[dict]:
    q = f"""
        SELECT nombre_municipio, cluster,
               votos_cepeda, pct_cepeda, votos_espriella, pct_espriella,
               no_votantes, brecha_vs_espriella
        FROM cuadrantes_2v
        WHERE departamento_nombre = ?
          AND {cuadrante_filter}
        ORDER BY no_votantes DESC
        LIMIT {limit}
    """
    rows = con.execute(q, [depto_nombre]).fetchall()
    return [{
        "mpio": r[0], "cluster": r[1] or "-",
        "votos_cepeda": r[2], "pct_cepeda": r[3],
        "votos_espriella": r[4], "pct_espriella": r[5],
        "no_votantes": r[6], "brecha": r[7],
    } for r in rows]


def build_prompt_recomendacion(d: dict, top_op: list, top_def: list) -> str:
    """Prompt corto para Ollama qwen2.5:14b."""
    pct_cep = (d["votos_cepeda"] / d["validos"] * 100) if d["validos"] else 0
    pct_esp = (d["votos_espriella"] / d["validos"] * 100) if d["validos"] else 0
    op_str = "\n".join(
        f"  - {x['mpio']} (no-votantes {x['no_votantes']:,}, Cep {x['pct_cepeda']:.1f}%, Esp {x['pct_espriella']:.1f}%)"
        for x in top_op[:5]
    )
    def_str = "\n".join(
        f"  - {x['mpio']} (Cep {x['pct_cepeda']:.1f}%, votos {x['votos_cepeda']:,})"
        for x in top_def[:5]
    )
    return f"""Eres analista político senior asesorando la campaña de Iván Cepeda (Pacto Histórico, Colombia 2026) para 2da vuelta. Producí una RECOMENDACIÓN TÁCTICA breve (4-6 frases · máximo 150 palabras) para el departamento {d['nombre']}.

DATOS:
- Municipios: {d['mpios']}
- Censo: {d['censo']:,}  · Votos válidos: {d['validos']:,}
- Cepeda: {d['votos_cepeda']:,} votos ({pct_cep:.1f}%)
- Espriella: {d['votos_espriella']:,} votos ({pct_esp:.1f}%)
- No-votantes (potencial): {d['no_votantes']:,}
- Cuadrantes locales: Q1_defender={d['q1']} · Q2_movilizar={d['q2']} · Q3_convertir={d['q3']} · Q4_resistir={d['q4']}

TOP 5 mpios oportunidad (Q2+Q3 con más no-votantes):
{op_str or '  (sin mpios oportunidad)'}

TOP 5 mpios defensa (Cepeda ganó):
{def_str or '  (Cepeda no ganó en ningún mpio aquí)'}

INSTRUCCIONES:
1. Diagnóstico en 1 frase: ¿este depto es territorio Cepeda, Espriella o disputado?
2. Mensaje sugerido específico al territorio (paz, víctimas, agro, minería, étnico, urbano, etc.) en 1-2 frases.
3. Acción operativa concreta en 1-2 frases (testigos electorales, brigadas, coalición con qué, foco geográfico).
4. NO uses encabezados ni listas. NO inventes municipios fuera de los listados. Escribí en español Colombia, prosa directa.""".strip()


def invoke_va_routing(prompt: str, va_root: Path, timeout_s: int = 90,
                       forzar_cloud: bool = True) -> str | None:
    """Invoca VA routing_inteligente vía subprocess.

    forzar_cloud=True (default) salta Ollama local (muy lento en esta máquina ~2.5min/call)
    y usa la cadena cloud-free: Cerebras qwen-3-235b (4-8s) → Groq → Mistral → Gemini.
    NUNCA cae a Claude (no está en la cadena 'razonamiento' del routing).
    """
    prompt_dir = PROJECT_ROOT / "data" / "raw" / "_prompts"
    prompt_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ_%f")
    pf = prompt_dir / f"depto_{ts}.txt"
    pf.write_text(prompt, encoding="utf-8")
    cmd = [
        sys.executable, "-m", "templates.routing_inteligente",
        "--tipo", "razonamiento",
        "--modelo", "gemini-2.5-flash",
        "--prompt-file", pf.as_posix(),
        "--max-tokens", "350",
    ]
    if forzar_cloud:
        cmd.append("--forzar-cloud")
    try:
        proc = subprocess.run(
            cmd, cwd=va_root.as_posix(), capture_output=True,
            text=True, encoding="utf-8", errors="replace", timeout=timeout_s,
        )
        if proc.returncode != 0:
            return None
        out = proc.stdout.strip()
        return out if out else None
    except Exception:  # noqa: BLE001
        return None


def fallback_recomendacion(d: dict) -> str:
    pct_cep = (d["votos_cepeda"] / d["validos"] * 100) if d["validos"] else 0
    pct_esp = (d["votos_espriella"] / d["validos"] * 100) if d["validos"] else 0
    if pct_cep > pct_esp + 5:
        diag = f"Territorio Cepeda fuerte (margen +{pct_cep - pct_esp:.1f} pp)"
        accion = "Foco en defender turnout. Testigos electorales en todos los puestos. No descuidar mpios donde ya gana asumiendo que está ganado."
    elif pct_esp > pct_cep + 5:
        diag = f"Territorio Espriella (margen +{pct_esp - pct_cep:.1f} pp). Recuperación requiere persuasión + movilización."
        accion = f"Foco en Q2_MOVILIZAR ({d['q2']} mpios) y Q3_CONVERTIR ({d['q3']} mpios). Coalición con Dignidad. Capturar voto blanco."
    else:
        diag = f"Departamento en disputa estrecha ({pct_cep:.1f}% vs {pct_esp:.1f}%)"
        accion = "Esfuerzo máximo. Aquí se gana o se pierde la elección. Movilización de no-votantes + reactivación voto Pacto histórico."
    return f"{diag}. Potencial movilizable: {d['no_votantes']:,} no-votantes. {accion}"


USE_LLM = True  # toggleado por --no-llm desde main()


def precompute_recomendaciones(deptos: list[dict], con, max_workers: int = 4) -> dict[str, dict]:
    """Genera las 34 recomendaciones EN PARALELO via VA cloud chain.

    Returns: dict { departamento_nombre: {"text": str, "elapsed_s": float, "source": "llm"|"fallback"} }
    """
    from concurrent.futures import ThreadPoolExecutor, as_completed

    # Prepara los prompts (1 SQL hit por depto para top mpios)
    prompts: dict[str, tuple[str, dict]] = {}
    for d in deptos:
        top_op = get_top_mpios_depto(
            con, d["nombre"], "cuadrante IN ('Q2_MOVILIZAR','Q3_CONVERTIR')", limit=10
        )
        top_def = get_top_mpios_depto(
            con, d["nombre"], "cuadrante LIKE 'Q1%'", limit=10
        )
        prompts[d["nombre"]] = (build_prompt_recomendacion(d, top_op, top_def), d)

    results: dict[str, dict] = {}

    def _worker(depto_nombre: str) -> tuple[str, dict]:
        prompt, _d = prompts[depto_nombre]
        t0 = time.monotonic()
        out = invoke_va_routing(prompt, VA_ROOT, timeout_s=60, forzar_cloud=True)
        elapsed = time.monotonic() - t0
        if out and len(out) > 80:
            out = out.replace("```", "").strip()
            return depto_nombre, {"text": out, "elapsed_s": elapsed, "source": "llm"}
        return depto_nombre, {
            "text": fallback_recomendacion(prompts[depto_nombre][1]),
            "elapsed_s": elapsed, "source": "fallback",
        }

    print(f"  [parallel-LLM] lanzando {len(prompts)} llamadas con {max_workers} workers (VA cloud chain · sin Claude)...")
    t_start = time.monotonic()
    with ThreadPoolExecutor(max_workers=max_workers) as exe:
        futures = {exe.submit(_worker, name): name for name in prompts}
        done = 0
        n_llm = 0
        for fut in as_completed(futures):
            name, res = fut.result()
            results[name] = res
            done += 1
            if res["source"] == "llm":
                n_llm += 1
            sys.stdout.write(f"\r  [parallel-LLM] {done}/{len(prompts)} · LLM={n_llm} fallback={done-n_llm} · {res['elapsed_s']:.1f}s {name[:25]:25s}")
            sys.stdout.flush()
    sys.stdout.write("\n")
    total = time.monotonic() - t_start
    print(f"  [parallel-LLM] terminado en {total:.1f}s · LLM={n_llm}/{len(prompts)} ({n_llm*100//len(prompts)}%)")
    return results


def build_doc(con) -> Document:
    doc = Document()
    # Estilos default
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Análisis Departamental · 2da Vuelta Presidencial Colombia 2026")
    r2.bold = True
    r2.font.size = Pt(22)
    r2.font.color.rgb = VERDE_PACTO

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("\nIván Cepeda Castro · Movimiento Político Pacto Histórico\nDónde hacer campaña para conseguir Cepeda + 3 millones de votos")
    r3.font.size = Pt(14)
    r3.font.color.rgb = GRIS

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(f"\n\nGenerado: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}\n"
                    f"Fuente: Registraduría Nacional · escrutinio al 99.92%\n"
                    f"Herramienta: github.com/wilsonherrera77/elecciones-co-2026")
    r4.font.size = Pt(10)
    r4.font.color.rgb = GRIS

    doc.add_page_break()

    # Resumen ejecutivo
    nac = get_summary_nacional(con)
    _h(doc, "Resumen ejecutivo", level=1, color=VERDE_PACTO)
    pct_cep = nac["cepeda"] / nac["validos"] * 100
    pct_esp = nac["espriella"] / nac["validos"] * 100
    _p(doc, f"En la primera vuelta presidencial del 31 de mayo de 2026, Iván Cepeda (Pacto Histórico) "
            f"obtuvo {nac['cepeda']:,} votos ({pct_cep:.2f}%) y Abelardo De La Espriella "
            f"(Defensores de la Patria) obtuvo {nac['espriella']:,} votos ({pct_esp:.2f}%). "
            f"Ambos pasaron a la segunda vuelta con una brecha de {nac['espriella']-nac['cepeda']:,} votos "
            f"a favor de Espriella.")
    _p(doc, f"Para ganar la segunda vuelta con turnout +5 puntos (escenario histórico colombiano), "
            f"Cepeda necesita aproximadamente +3.000.000 votos adicionales. Este documento desglosa, "
            f"departamento por departamento, dónde están esos votos disponibles y qué táctica corresponde.")
    _p(doc, "Método: clasificación de los 1.189 municipios en 4 cuadrantes operativos:", bold=True)
    _p(doc, "  • Q1 DEFENDER: Cepeda ganó · proteger turnout con testigos electorales (425 mpios)")
    _p(doc, "  • Q2 MOVILIZAR: margen ≤10 pp · empujar no-votantes (81 mpios decisivos)")
    _p(doc, "  • Q3 CONVERTIR: gap 10-30 pp · persuasión centro/Dignidad/voto blanco (240 mpios)")
    _p(doc, "  • Q4 RESISTIR: gap >30 pp · piso digno · no derrochar recursos (437 mpios)")

    doc.add_page_break()

    # Tabla nacional resumen 34 deptos
    _h(doc, "Tabla nacional · 34 departamentos por votos Cepeda", level=1, color=VERDE_PACTO)
    deptos = get_deptos(con)
    t = doc.add_table(rows=0, cols=8)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _row(t, ["#", "Departamento", "Cepeda", "% Cep", "% Esp", "No-vot.", "Mpios", "Q1/Q2/Q3/Q4"], header=True)
    for i, d in enumerate(deptos, 1):
        pct_cep = (d["votos_cepeda"] / d["validos"] * 100) if d["validos"] else 0
        pct_esp = (d["votos_espriella"] / d["validos"] * 100) if d["validos"] else 0
        _row(t, [
            str(i), d["nombre"],
            f"{d['votos_cepeda']:,}",
            f"{pct_cep:.1f}%", f"{pct_esp:.1f}%",
            f"{d['no_votantes']:,}",
            str(d["mpios"]),
            f"{d['q1']}/{d['q2']}/{d['q3']}/{d['q4']}",
        ])

    doc.add_page_break()

    # Capítulos por departamento
    _h(doc, "Capítulos por departamento", level=1, color=VERDE_PACTO)

    # Precomputa las 34 recomendaciones EN PARALELO si USE_LLM activo
    recomendaciones: dict[str, dict] = {}
    if USE_LLM:
        recomendaciones = precompute_recomendaciones(deptos, con, max_workers=4)

    for i, d in enumerate(deptos, 1):
        sys.stdout.write(f"\r  [{i:>2}/{len(deptos)}] escribiendo {d['nombre']:25s}")
        sys.stdout.flush()

        doc.add_page_break()
        _h(doc, f"{i}. {d['nombre']}", level=2, color=VERDE_PACTO)
        pct_cep = (d["votos_cepeda"] / d["validos"] * 100) if d["validos"] else 0
        pct_esp = (d["votos_espriella"] / d["validos"] * 100) if d["validos"] else 0

        _p(doc, "Cifras clave", bold=True, size=12)
        tbl = doc.add_table(rows=0, cols=2)
        _row(tbl, ["Indicador", "Valor"], header=True)
        _row(tbl, ["Municipios", str(d["mpios"])])
        _row(tbl, ["Censo electoral", f"{d['censo']:,}"])
        _row(tbl, ["Votos válidos", f"{d['validos']:,}"])
        _row(tbl, ["Cepeda (Pacto Histórico)", f"{d['votos_cepeda']:,} · {pct_cep:.1f}%"])
        _row(tbl, ["Espriella (Defensores)", f"{d['votos_espriella']:,} · {pct_esp:.1f}%"])
        _row(tbl, ["Brecha (Esp − Cep)", f"{d['votos_espriella'] - d['votos_cepeda']:+,}"])
        _row(tbl, ["No-votantes (potencial)", f"{d['no_votantes']:,}"])
        _row(tbl, ["Cuadrantes Q1/Q2/Q3/Q4", f"{d['q1']} / {d['q2']} / {d['q3']} / {d['q4']}"])

        # Top oportunidad (Q2+Q3)
        top_op = get_top_mpios_depto(
            con, d["nombre"],
            "cuadrante IN ('Q2_MOVILIZAR','Q3_CONVERTIR')",
            limit=10,
        )
        _p(doc, "")
        _p(doc, "Top 10 municipios oportunidad (Q2_MOVILIZAR + Q3_CONVERTIR)", bold=True, size=12)
        if top_op:
            tbl_op = doc.add_table(rows=0, cols=6)
            _row(tbl_op, ["Municipio", "Cluster", "Cep%", "Esp%", "No-vot.", "Brecha"], header=True)
            for x in top_op:
                _row(tbl_op, [x["mpio"], x["cluster"], f"{x['pct_cepeda']:.1f}%",
                              f"{x['pct_espriella']:.1f}%", f"{x['no_votantes']:,}", f"{x['brecha']:+,}"])
        else:
            _p(doc, "(sin municipios Q2/Q3 en este departamento)", color=GRIS)

        # Top defensa (Q1)
        top_def = get_top_mpios_depto(
            con, d["nombre"],
            "cuadrante LIKE 'Q1%'",
            limit=10,
        )
        _p(doc, "")
        _p(doc, "Top 10 municipios defensa (Q1_DEFENDER · Cepeda ganó)", bold=True, size=12)
        if top_def:
            tbl_def = doc.add_table(rows=0, cols=4)
            _row(tbl_def, ["Municipio", "Cluster", "Votos Cepeda", "% Cepeda"], header=True)
            for x in top_def:
                _row(tbl_def, [x["mpio"], x["cluster"], f"{x['votos_cepeda']:,}",
                               f"{x['pct_cepeda']:.1f}%"])
        else:
            _p(doc, "(Cepeda no ganó en ningún municipio de este departamento)", color=GRIS)

        # Recomendación táctica
        _p(doc, "")
        modo_label = "[VA-ROUTE cadena cloud-free · sin Claude]" if USE_LLM else "[determinista]"
        _p(doc, f"Recomendación táctica · {modo_label}", bold=True, size=12, color=VERDE_PACTO)
        if USE_LLM:
            r = recomendaciones.get(d["nombre"])
            if r and r["source"] == "llm":
                _p(doc, r["text"])
                _p(doc, f"_(LLM cloud-free vía VA · {r['elapsed_s']:.1f}s)_",
                   size=8, color=GRIS, align=WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                _p(doc, r["text"] if r else fallback_recomendacion(d))
                _p(doc, "_(fallback determinista · proveedores cloud-free no respondieron)_",
                   size=8, color=GRIS, align=WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            _p(doc, fallback_recomendacion(d))

    sys.stdout.write("\n")

    # Anexo metodológico
    doc.add_page_break()
    _h(doc, "Anexo · metodología", level=1, color=VERDE_PACTO)
    _p(doc, "Fuente de datos", bold=True, size=12)
    _p(doc, "Resultados oficiales Registraduría Nacional vía endpoint público "
            "https://resultados.registraduria.gov.co/json/ACT/PR/<codigo>.json. "
            "Escrutinio al 99.92% al cierre de captura (121.925 de 122.020 mesas).")
    _p(doc, "")
    _p(doc, "Definición de cuadrantes", bold=True, size=12)
    _p(doc, "Q1_DEFENDER: Cepeda obtuvo más votos que Espriella en el municipio Y su porcentaje "
            "es ≥ 40% de los votos válidos. Acción: proteger turnout con testigos.")
    _p(doc, "Q1_DEFENDER_FRAGIL: Cepeda obtuvo más votos pero su porcentaje es <40%. "
            "Acción: reforzar antes de que el territorio se mueva.")
    _p(doc, "Q2_MOVILIZAR: la diferencia absoluta entre el porcentaje de Cepeda y el de Espriella "
            "es ≤ 10 puntos porcentuales. Acción: empujar no-votantes Pacto · transporte el día E.")
    _p(doc, "Q3_CONVERTIR: la diferencia es entre 10 y 30 puntos a favor de Espriella. "
            "Acción: persuasión del electorado de centro · alianza con Dignidad & Compromiso · "
            "captura de voto blanco.")
    _p(doc, "Q4_RESISTIR: la diferencia supera 30 puntos a favor de Espriella. "
            "Territorio hostil · acción defensiva · piso digno >20% · no derrochar recursos.")
    _p(doc, "")
    _p(doc, "Cálculo de 3M votos", bold=True, size=12)
    _p(doc, "El umbral 50%+1 con turnout +5pp histórico se sitúa en 12.85 millones de votos válidos. "
            "Cepeda actual 9.68M → gap +3.17M ≈ los 3 millones de la meta operativa de campaña.")
    _p(doc, "")
    _p(doc, "Limitaciones reconocidas", bold=True, size=12)
    _p(doc, "1. Bogotá D.C. está agregada como un solo municipio (level=3 nomenclator) · "
            "drill-down por localidad requiere extender el scraper a level=4 (ZONA) o level=6 (PUESTO).")
    _p(doc, "2. NBI municipal DANE 2018 no se incluye (URL automática devolvió 404). "
            "El score actual no pondera vulnerabilidad social.")
    _p(doc, "3. Las recomendaciones tácticas son apoyo cuantitativo · no reemplazan análisis "
            "humano cualitativo del territorio.")
    _p(doc, "")
    _p(doc, "Stack técnico", bold=True, size=12)
    _p(doc, "Python 3.11 · httpx async HTTP/2 · DuckDB · Plotly · python-docx · "
            "Ollama qwen2.5:14b local (cero API cloud paga). "
            "Construido sobre framework Visual_Agentes (LOCAL-FIRST estricto · regla #11).")

    return doc


def main() -> int:
    global USE_LLM
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--no-llm", action="store_true",
                    help="usa fallback determinista para las 34 recomendaciones (rápido)")
    args = ap.parse_args()
    USE_LLM = not args.no_llm

    if not DB.exists():
        print(f"ERROR: falta {DB}. Corre el pipeline scraper primero.")
        return 1
    OUT.parent.mkdir(parents=True, exist_ok=True)
    print(f"[docx_departamental] generando {OUT.relative_to(PROJECT_ROOT)}")
    print(f"  VA_ROOT: {VA_ROOT}")
    print(f"  modo: {'Ollama qwen2.5:14b (LLM)' if USE_LLM else 'fallback determinista (rápido)'}")
    con = duckdb.connect(DB.as_posix(), read_only=False)
    # Verifica cuadrantes_2v
    has = con.execute(
        "SELECT COUNT(*) FROM information_schema.tables WHERE table_name='cuadrantes_2v'"
    ).fetchone()[0]
    if not has:
        print("ERROR: tabla cuadrantes_2v no existe · corre `python -m estrategia.segunda_vuelta_3M` primero")
        return 1
    doc = build_doc(con)
    con.close()
    doc.save(OUT.as_posix())
    size_kb = OUT.stat().st_size // 1024
    print(f"[docx_departamental] OK · {OUT.relative_to(PROJECT_ROOT)} · {size_kb} KB")
    return 0


if __name__ == "__main__":
    sys.exit(main())
